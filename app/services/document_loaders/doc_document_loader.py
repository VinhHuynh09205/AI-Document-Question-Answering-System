import re
import shutil
import subprocess
from pathlib import Path

from langchain_core.documents import Document

from app.services.interfaces.document_loader import IDocumentLoader


class DocDocumentLoader(IDocumentLoader):
    """Loader for legacy .doc (Word 97-2003) files."""

    def supports(self, file_extension: str) -> bool:
        return file_extension.lower() == ".doc"

    def load(self, file_path: Path) -> list[Document]:
        content = self._try_as_docx(file_path)
        if not content.strip():
            content = self._try_antiword(file_path)
        if not content.strip():
            content = self._extract_text_binary(file_path)

        return [
            Document(
                page_content=content,
                metadata={"source": str(file_path), "extension": ".doc"},
            )
        ]

    @staticmethod
    def _try_as_docx(file_path: Path) -> str:
        """Some .doc files are actually .docx (Office Open XML) with wrong extension."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(file_path))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            table_lines: list[str] = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        table_lines.append(" | ".join(cells))
            return "\n".join(paragraphs + table_lines)
        except Exception:
            return ""

    @staticmethod
    def _try_antiword(file_path: Path) -> str:
        """Use antiword CLI tool if available (installed in Docker)."""
        if shutil.which("antiword") is None:
            return ""
        try:
            result = subprocess.run(
                ["antiword", "-m", "UTF-8", str(file_path)],
                capture_output=True,
                text=True,
                timeout=30,
            )
            return result.stdout.strip() if result.returncode == 0 else ""
        except Exception:
            return ""

    @staticmethod
    def _extract_text_binary(file_path: Path) -> str:
        """Best-effort text extraction from binary .doc format."""
        with open(file_path, "rb") as f:
            data = f.read()

        # Word 97-2003 stores text in cp1252 encoding
        text = data.decode("cp1252", errors="replace")

        # Remove null bytes and binary control characters
        text = re.sub(r"[\x00]+", " ", text)
        text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]+", " ", text)

        lines: list[str] = []
        for segment in text.split("\r"):
            cleaned = segment.strip()
            if not cleaned or len(cleaned) < 4:
                continue
            # Keep lines with majority readable characters
            readable = sum(1 for c in cleaned if c.isalnum() or c.isspace() or c in ".,;:!?()-")
            if readable > len(cleaned) * 0.6:
                lines.append(cleaned)

        return "\n".join(lines)
