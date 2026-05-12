from pathlib import Path
import hashlib
import re
import time
from zipfile import ZipFile

from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from langchain_core.documents import Document

from app.services.interfaces.document_loader import IDocumentLoader
from app.services.interfaces.image_understanding_service import IImageUnderstandingService


class DocxDocumentLoader(IDocumentLoader):
    _OCR_MARKER_RE = re.compile(
        r"(local_ocr|local_vision|image\s*analysis|slide\s*image|provider[:=])",
        re.IGNORECASE,
    )
    _LATIN_WORD_RE = re.compile(r"[A-Za-z\u00C0-\u024F]{2,}")
    _CJK_RE = re.compile(r"[\u3040-\u30FF\u3400-\u4DBF\u4E00-\u9FFF]")
    _NUMERIC_TOKEN_RE = re.compile(r"\d+(?:[.,]\d+)?")
    _HEADING_LEVEL_RE = re.compile(r"heading\s*(\d+)", re.IGNORECASE)

    def __init__(
        self,
        image_understanding_service: IImageUnderstandingService | None = None,
        *,
        max_images_per_document: int = 4,
        text_char_threshold_for_image_analysis: int = 2200,
        max_image_analysis_seconds: float = 20.0,
    ) -> None:
        self._image_understanding_service = image_understanding_service
        self._max_images_per_document = max(1, max_images_per_document)
        self._text_char_threshold_for_image_analysis = max(300, text_char_threshold_for_image_analysis)
        self._max_image_analysis_seconds = max(1.0, float(max_image_analysis_seconds))

    def supports(self, file_extension: str) -> bool:
        return file_extension.lower() == ".docx"

    def load(self, file_path: Path) -> list[Document]:
        doc = DocxDocument(str(file_path))

        page_contents = self._extract_docx_page_contents(doc)
        body_text_snapshot = "\n\n".join(content for _, content in page_contents)
        image_notes, providers_used = self._extract_docx_image_notes(file_path, body_text_snapshot)

        if image_notes:
            if page_contents:
                last_page_number, last_content = page_contents[-1]
                merged_content = (
                    f"{last_content}\n\n[Image insights]\n" + "\n".join(image_notes)
                ).strip()
                page_contents[-1] = (last_page_number, merged_content)
            else:
                notes_block = "\n".join(image_notes)
                page_contents = [
                    (1, f"[Image insights]\n{notes_block}".strip())
                ]

        if not page_contents:
            return []

        total_pages = max(page_number for page_number, _ in page_contents)
        ocr_applied = "local_ocr" in providers_used
        return [
            Document(
                page_content=content,
                metadata={
                    "source": str(file_path),
                    "extension": ".docx",
                    "page": page_number,
                    "total_pages": total_pages,
                    "content_type": "docx_page",
                    "image_analysis_applied": bool(image_notes),
                    "ocr_applied": ocr_applied,
                },
            )
            for page_number, content in page_contents
        ]

    @classmethod
    def _extract_docx_page_contents(cls, doc: DocxDocument) -> list[tuple[int, str]]:
        paragraph_by_element = {id(paragraph._element): paragraph for paragraph in doc.paragraphs}
        table_by_element = {id(table._element): table for table in doc.tables}

        page_lines: dict[int, list[str]] = {1: []}
        current_page = 1
        seen: set[str] = set()

        for child in doc.element.body.iterchildren():
            child_tag = str(getattr(child, "tag", "")).rsplit("}", 1)[-1]

            if child_tag == "p":
                paragraph = paragraph_by_element.get(id(child))
                if paragraph is None:
                    continue

                line = cls._format_paragraph_line(paragraph, seen)
                if line:
                    page_lines.setdefault(current_page, []).append(line)

                if cls._paragraph_has_page_break(paragraph):
                    current_page += 1
                    page_lines.setdefault(current_page, [])
                continue

            if child_tag == "tbl":
                table = table_by_element.get(id(child))
                if table is None:
                    continue
                table_lines = cls._extract_table_lines(table)
                if table_lines:
                    page_lines.setdefault(current_page, []).extend(table_lines)

        page_contents: list[tuple[int, str]] = []
        for page_number in sorted(page_lines):
            content = "\n".join(page_lines[page_number]).strip()
            if content:
                page_contents.append((page_number, content))

        if page_contents:
            return page_contents

        fallback_lines: list[str] = []
        fallback_seen: set[str] = set()
        for paragraph in doc.paragraphs:
            line = cls._format_paragraph_line(paragraph, fallback_seen)
            if line:
                fallback_lines.append(line)
        for table in doc.tables:
            fallback_lines.extend(cls._extract_table_lines(table))

        fallback_content = "\n".join(fallback_lines).strip()
        if fallback_content:
            return [(1, fallback_content)]
        return []

    @classmethod
    def _format_paragraph_line(cls, paragraph: DocxParagraph, seen: set[str]) -> str:
        raw_text = str(paragraph.text or "")
        text = re.sub(r"\s+", " ", raw_text).strip()
        if not text:
            return ""

        style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "").strip()
        normalized_style = style_name.lower()

        heading_match = cls._HEADING_LEVEL_RE.search(normalized_style)
        if heading_match:
            level = max(1, min(6, int(heading_match.group(1))))
            line = f"{'#' * level} {text}"
        elif "list" in normalized_style or text.startswith(("- ", "* ", "• ")):
            clean_text = text.lstrip("-*• ").strip()
            line = f"- {clean_text}" if clean_text else ""
        else:
            line = text

        if not line:
            return ""

        dedupe_key = line.lower()
        if dedupe_key in seen:
            return ""
        seen.add(dedupe_key)
        return line

    @staticmethod
    def _paragraph_has_page_break(paragraph: DocxParagraph) -> bool:
        paragraph_xml = str(paragraph._element.xml or "")
        return (
            "lastRenderedPageBreak" in paragraph_xml
            or 'w:type="page"' in paragraph_xml
            or "w:type='page'" in paragraph_xml
        )

    @staticmethod
    def _extract_table_lines(table: DocxTable) -> list[str]:
        lines: list[str] = []
        for row in table.rows:
            cells = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
        return lines

    def _extract_docx_image_notes(
        self,
        file_path: Path,
        text_snapshot: str,
    ) -> tuple[list[str], set[str]]:
        if self._image_understanding_service is None:
            return [], set()

        compact_text = str(text_snapshot or "").strip()
        if len(compact_text) >= self._text_char_threshold_for_image_analysis:
            return [], set()

        started_at = time.perf_counter()
        notes: list[str] = []
        providers_used: set[str] = set()
        seen_image_hashes: set[str] = set()
        doc_prefers_cjk = self._contains_substantial_cjk(text_snapshot)

        try:
            with ZipFile(file_path, "r") as archive:
                media_entries = [
                    name for name in archive.namelist()
                    if name.startswith("word/media/") and not name.endswith("/")
                ]
                for image_index, entry_name in enumerate(media_entries, start=1):
                    if image_index > self._max_images_per_document:
                        break
                    if (time.perf_counter() - started_at) >= self._max_image_analysis_seconds:
                        break

                    image_bytes = archive.read(entry_name)
                    if not image_bytes:
                        continue

                    image_hash = self._fingerprint_image_bytes(image_bytes)
                    if image_hash in seen_image_hashes:
                        continue
                    seen_image_hashes.add(image_hash)

                    result = self._image_understanding_service.analyze_image(
                        image_bytes,
                        source=str(file_path),
                        hint=f"docx image {image_index}",
                    )
                    providers_used.add(str(result.provider or "").strip().lower())
                    note_text = self._normalize_image_note_text(str(result.text or ""))
                    if not note_text:
                        continue

                    if not self._is_useful_image_note(
                        note_text,
                        document_prefers_cjk=doc_prefers_cjk,
                    ):
                        continue

                    if self._is_duplicate_image_note(
                        note_text,
                        text_snapshot,
                        notes,
                    ):
                        continue

                    notes.append(f"Image {image_index}: {note_text}")
        except OSError:
            return [], set()

        return notes, providers_used

    @staticmethod
    def _fingerprint_image_bytes(image_bytes: bytes) -> str:
        return hashlib.sha1(image_bytes).hexdigest()[:24]

    @classmethod
    def _normalize_image_note_text(cls, text: str) -> str:
        normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
        cleaned_lines: list[str] = []
        seen: set[str] = set()

        for raw_line in normalized.splitlines():
            line = re.sub(r"\s+", " ", raw_line).strip(" -•\t")
            if not line:
                continue
            if cls._OCR_MARKER_RE.search(line):
                continue

            lowered = line.lower()
            if lowered in seen:
                continue
            seen.add(lowered)
            cleaned_lines.append(line)
            if len(cleaned_lines) >= 5:
                break

        return "\n".join(cleaned_lines).strip()

    @classmethod
    def _contains_substantial_cjk(cls, text: str) -> bool:
        return len(cls._CJK_RE.findall(str(text or ""))) >= 6

    @classmethod
    def _is_useful_image_note(
        cls,
        note_text: str,
        *,
        document_prefers_cjk: bool,
    ) -> bool:
        compact = " ".join(str(note_text or "").split()).strip()
        if len(compact) < 8:
            return False

        if cls._OCR_MARKER_RE.search(compact):
            return False

        has_cjk = bool(cls._CJK_RE.search(compact))
        if document_prefers_cjk and not has_cjk:
            return False

        stripped_symbols = re.sub(
            r"[A-Za-z\u00C0-\u024F0-9\u3040-\u30FF\u3400-\u4DBF\u4E00-\u9FFF\s]",
            "",
            compact,
        )
        if len(stripped_symbols) > len(compact) * 0.35:
            return False

        if not has_cjk:
            latin_word_count = len(cls._LATIN_WORD_RE.findall(compact))
            if latin_word_count < 2:
                numeric_count = len(cls._NUMERIC_TOKEN_RE.findall(compact))
                if not (
                    (latin_word_count >= 1 and numeric_count >= 3)
                    or numeric_count >= 5
                    or "%" in compact
                ):
                    return False

        return True

    @staticmethod
    def _is_duplicate_image_note(
        note_text: str,
        text_snapshot: str,
        existing_notes: list[str],
    ) -> bool:
        compact_note = re.sub(r"\s+", " ", str(note_text or "")).strip().lower()
        if not compact_note:
            return True

        compact_doc = re.sub(r"\s+", " ", str(text_snapshot or "")).strip().lower()
        if compact_note in compact_doc:
            return True

        for existing in existing_notes:
            _, _, existing_note = existing.partition(":")
            normalized_existing = re.sub(r"\s+", " ", existing_note).strip().lower()
            if compact_note == normalized_existing:
                return True

        return False
