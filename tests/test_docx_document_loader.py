from docx import Document as DocxDocument
from docx.enum.text import WD_BREAK

from app.services.document_loaders.docx_document_loader import DocxDocumentLoader


def test_docx_loader_sets_page_metadata_with_page_break(tmp_path) -> None:
    file_path = tmp_path / "sample.docx"

    doc = DocxDocument()
    doc.add_paragraph("Noi dung o trang 1")
    break_paragraph = doc.add_paragraph()
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("Noi dung o trang 2")
    doc.save(file_path)

    loader = DocxDocumentLoader()
    docs = loader.load(file_path)

    assert len(docs) == 2
    assert docs[0].metadata.get("page") == 1
    assert docs[1].metadata.get("page") == 2
    assert "Noi dung o trang 1" in docs[0].page_content
    assert "Noi dung o trang 2" in docs[1].page_content


def test_docx_loader_defaults_to_single_page_when_no_break(tmp_path) -> None:
    file_path = tmp_path / "single-page.docx"

    doc = DocxDocument()
    doc.add_paragraph("Noi dung khong co page break")
    doc.save(file_path)

    loader = DocxDocumentLoader()
    docs = loader.load(file_path)

    assert len(docs) == 1
    assert docs[0].metadata.get("page") == 1
    assert docs[0].metadata.get("total_pages") == 1


def test_docx_loader_accepts_numeric_chart_note_without_two_latin_words() -> None:
    assert (
        DocxDocumentLoader._is_useful_image_note(
            "trieungoi 26.9 24.2 22.3 20.8",
            document_prefers_cjk=False,
        )
        is True
    )
